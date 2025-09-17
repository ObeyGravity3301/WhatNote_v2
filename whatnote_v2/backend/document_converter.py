"""
文档转换服务
支持Word文档转换为PDF
"""
import os
import tempfile
from pathlib import Path
from typing import Optional
import pypandoc
from docx import Document
import subprocess
import shutil
import docx2txt
import win32com.client
import time

class DocumentConverter:
    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "whatnote_converter"
        self.temp_dir.mkdir(exist_ok=True)
    
    def convert_word_to_pdf(self, word_file_path: str, output_dir: str) -> Optional[str]:
        """
        将Word文档转换为PDF
        
        Args:
            word_file_path: Word文档路径
            output_dir: 输出目录
            
        Returns:
            转换后的PDF文件路径，失败返回None
        """
        try:
            word_path = Path(word_file_path)
            if not word_path.exists():
                print(f"Word文件不存在: {word_file_path}")
                return None
            
            # 生成输出PDF文件名
            pdf_filename = word_path.stem + ".pdf"
            pdf_path = Path(output_dir) / pdf_filename
            
            print(f"开始转换Word文档: {word_path.name} -> {pdf_filename}")
            
            # 方法1: 使用Microsoft Office COM接口转换（最高质量）
            try:
                result = self._convert_with_office_com(word_path, pdf_path)
                if result and result.endswith('.pdf') and Path(result).exists():
                    print(f"Office COM转换成功: {pdf_path}")
                    return result
            except Exception as e:
                print(f"Office COM转换失败: {e}")
            
            # 方法2: 使用系统打印驱动转换
            try:
                result = self._convert_with_print_driver(word_path, pdf_path)
                if result and result.endswith('.pdf') and Path(result).exists():
                    print(f"打印驱动转换成功: {pdf_path}")
                    return result
            except Exception as e:
                print(f"打印驱动转换失败: {e}")
            
            # 方法3: 使用pypandoc转换
            try:
                pypandoc.convert_file(
                    str(word_path),
                    'pdf',
                    outputfile=str(pdf_path),
                    extra_args=['--pdf-engine=xelatex']
                )
                if pdf_path.exists():
                    print(f"pypandoc转换成功: {pdf_path}")
                    return str(pdf_path)
            except Exception as e:
                print(f"pypandoc转换失败: {e}")
            
            # 方法4: 使用LibreOffice转换
            try:
                result = self._convert_with_libreoffice(word_path, pdf_path)
                if result and result.endswith('.pdf') and Path(result).exists():
                    print(f"LibreOffice转换成功: {pdf_path}")
                    return result
            except Exception as e:
                print(f"LibreOffice转换失败: {e}")
            
            print(f"所有PDF转换方法都失败了: {word_path.name}")
            return None
            
        except Exception as e:
            print(f"Word转PDF转换异常: {e}")
            return None
    
    def _convert_with_libreoffice(self, word_path: Path, pdf_path: Path) -> Optional[str]:
        """使用LibreOffice转换"""
        try:
            # 检查LibreOffice是否安装
            result = subprocess.run(['libreoffice', '--version'], 
                                 capture_output=True, text=True, timeout=10)
            if result.returncode != 0:
                raise Exception("LibreOffice未安装")
            
            # 使用LibreOffice转换
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(pdf_path.parent),
                str(word_path)
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            if result.returncode == 0 and pdf_path.exists():
                print(f"LibreOffice转换成功: {pdf_path}")
                return str(pdf_path)
            else:
                raise Exception(f"LibreOffice转换失败: {result.stderr}")
                
        except Exception as e:
            print(f"LibreOffice转换异常: {e}")
            raise
    
    def _convert_with_docx_to_html(self, word_path: Path, pdf_path: Path) -> Optional[str]:
        """使用python-docx转换为HTML，然后尝试转PDF"""
        try:
            # 读取Word文档
            doc = Document(word_path)
            
            # 创建HTML内容
            html_content = self._docx_to_html(doc)
            
            # 保存为HTML文件
            html_path = word_path.with_suffix('.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # 尝试使用pypandoc将HTML转PDF
            try:
                pypandoc.convert_file(
                    str(html_path),
                    'pdf',
                    outputfile=str(pdf_path),
                    extra_args=['--pdf-engine=wkhtmltopdf']
                )
                
                # 清理HTML文件
                html_path.unlink()
                
                if pdf_path.exists():
                    print(f"HTML转PDF成功: {pdf_path}")
                    return str(pdf_path)
            except Exception as e:
                print(f"HTML转PDF失败: {e}")
            
            # 如果PDF转换失败，返回None（不返回HTML文件）
            print(f"HTML转PDF失败，清理HTML文件: {html_path}")
            html_path.unlink()  # 删除HTML文件
            return None
            
        except Exception as e:
            print(f"docx转HTML异常: {e}")
            raise
    
    def _convert_to_text(self, word_path: Path, output_dir: str) -> Optional[str]:
        """将Word文档转换为纯文本"""
        try:
            # 使用docx2txt提取文本
            text_content = docx2txt.process(str(word_path))
            
            if not text_content or not text_content.strip():
                print(f"Word文档为空或无法提取文本: {word_path.name}")
                return None
            
            # 生成文本文件名
            text_filename = word_path.stem + ".txt"
            text_path = Path(output_dir) / text_filename
            
            # 保存文本文件
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            print(f"Word文档转换为文本成功: {word_path.name} -> {text_filename}")
            return str(text_path)
            
        except Exception as e:
            print(f"转换为文本异常: {e}")
            raise
    
    def _convert_to_html_only(self, word_path: Path, output_dir: str) -> Optional[str]:
        """将Word文档转换为HTML文件（不转PDF）"""
        try:
            print(f"尝试转换为HTML: {word_path.name}")
            
            # 读取Word文档
            doc = Document(word_path)
            
            # 创建HTML内容
            html_content = self._docx_to_html(doc)
            
            # 生成HTML文件名
            html_filename = word_path.stem + ".html"
            html_path = Path(output_dir) / html_filename
            
            # 保存HTML文件
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            print(f"Word文档转换为HTML成功: {word_path.name} -> {html_filename}")
            return str(html_path)
            
        except Exception as e:
            print(f"转换为HTML异常: {e}")
            raise
    
    def _convert_with_office_com(self, word_path: Path, pdf_path: Path) -> Optional[str]:
        """使用Microsoft Office COM接口转换（最高质量）"""
        try:
            print(f"尝试使用Office COM接口转换: {word_path.name}")
            
            # 启动Word应用程序
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False  # 不显示Word窗口
            word_app.DisplayAlerts = False  # 不显示警告
            
            try:
                # 打开Word文档
                doc = word_app.Documents.Open(str(word_path.absolute()))
                
                # 导出为PDF
                doc.ExportAsFixedFormat(
                    OutputFileName=str(pdf_path.absolute()),
                    ExportFormat=17,  # PDF格式
                    OpenAfterExport=False,
                    OptimizeFor=0,  # 打印质量
                    BitmapMissingFonts=True,
                    DocStructureTags=True,
                    CreateBookmarks=0  # 创建书签
                )
                
                # 关闭文档
                doc.Close()
                
                if pdf_path.exists():
                    print(f"Office COM转换成功: {pdf_path}")
                    return str(pdf_path)
                else:
                    raise Exception("PDF文件未生成")
                    
            finally:
                # 关闭Word应用程序
                word_app.Quit()
                time.sleep(1)  # 等待Word完全关闭
                
        except Exception as e:
            print(f"Office COM转换异常: {e}")
            raise
    
    def _convert_with_print_driver(self, word_path: Path, pdf_path: Path) -> Optional[str]:
        """使用系统打印驱动转换"""
        try:
            print(f"尝试使用打印驱动转换: {word_path.name}")
            
            # 使用PowerShell调用Word的导出功能
            ps_script = f'''
            $word = New-Object -ComObject Word.Application
            $word.Visible = $false
            $word.DisplayAlerts = $false
            
            try {{
                $doc = $word.Documents.Open("{word_path.absolute()}")
                $doc.ExportAsFixedFormat("{pdf_path.absolute()}", 17, $false, 0, $true, $true, 0, $true)
                $doc.Close()
                Write-Output "转换成功"
            }} finally {{
                $word.Quit()
                [System.Runtime.Interopservices.Marshal]::ReleaseComObject($word) | Out-Null
            }}
            '''
            
            # 执行PowerShell脚本
            result = subprocess.run([
                'powershell', '-Command', ps_script
            ], capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0 and pdf_path.exists():
                print(f"打印驱动转换成功: {pdf_path}")
                return str(pdf_path)
            else:
                raise Exception(f"PowerShell转换失败: {result.stderr}")
                
        except Exception as e:
            print(f"打印驱动转换异常: {e}")
            raise
    
    def _docx_to_html(self, doc: Document) -> str:
        """将docx文档转换为HTML"""
        html_parts = []
        html_parts.append("""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Word Document</title>
    <style>
        body { font-family: 'Times New Roman', serif; margin: 40px; line-height: 1.6; }
        h1, h2, h3, h4, h5, h6 { color: #333; margin-top: 20px; margin-bottom: 10px; }
        p { margin-bottom: 10px; }
        .page-break { page-break-before: always; }
        table { border-collapse: collapse; width: 100%; margin: 10px 0; }
        td, th { border: 1px solid #ddd; padding: 8px; text-align: left; }
        th { background-color: #f2f2f2; }
    </style>
</head>
<body>
""")
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # 简单的段落样式检测
                style_name = paragraph.style.name.lower()
                if 'heading' in style_name:
                    level = 1
                    if 'heading 2' in style_name:
                        level = 2
                    elif 'heading 3' in style_name:
                        level = 3
                    elif 'heading 4' in style_name:
                        level = 4
                    elif 'heading 5' in style_name:
                        level = 5
                    elif 'heading 6' in style_name:
                        level = 6
                    
                    html_parts.append(f"<h{level}>{paragraph.text}</h{level}>")
                else:
                    html_parts.append(f"<p>{paragraph.text}</p>")
        
        # 处理表格
        for table in doc.tables:
            html_parts.append("<table>")
            for i, row in enumerate(table.rows):
                html_parts.append("<tr>")
                for cell in row.cells:
                    tag = "th" if i == 0 else "td"
                    html_parts.append(f"<{tag}>{cell.text}</{tag}>")
                html_parts.append("</tr>")
            html_parts.append("</table>")
        
        html_parts.append("</body></html>")
        return "\n".join(html_parts)
    
    def cleanup_temp_files(self):
        """清理临时文件"""
        try:
            if self.temp_dir.exists():
                shutil.rmtree(self.temp_dir)
                print("临时文件清理完成")
        except Exception as e:
            print(f"清理临时文件失败: {e}")

# 全局转换器实例
document_converter = DocumentConverter()
